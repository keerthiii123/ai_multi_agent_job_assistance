from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document


# =====================================================
# SAVE PDF REPORT
# =====================================================

def save_pdf(report, filename="job_readiness_report.pdf"):

    doc = SimpleDocTemplate(filename)

    styles = getSampleStyleSheet()

    story = []

    for line in report.split("\n"):

        story.append(
            Paragraph(line.replace("&", "&amp;"), styles["BodyText"])
        )

    doc.build(story)

    return filename


# =====================================================
# SAVE DOCX REPORT
# =====================================================

def save_docx(report, filename="job_readiness_report.docx"):

    document = Document()

    document.add_heading(
        "AI Multi-Agent Job Readiness Report",
        level=1
    )

    for line in report.split("\n"):

        document.add_paragraph(line)

    document.save(filename)

    return filename